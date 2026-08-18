from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "双Agent自然语言证明审计项目_进展与下一阶段规划书_2026-08-18.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "1F5E3B"
GOLD = "7A5A00"
RED = "9B1C1C"
MUTED = "5F6B76"


def set_run_font(run, east="Microsoft YaHei", west="Calibri", size=None, bold=None, color=None):
    run.font.name = west
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), west)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), west)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            color = None
            if status_col == idx:
                if "已完成" in str(value) or "通过" in str(value):
                    color = GREEN
                elif "阻塞" in str(value) or "未完成" in str(value):
                    color = RED
                else:
                    color = GOLD
            set_run_font(r, size=9.2, bold=(status_col == idx), color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        a = p.add_run(bold_lead)
        set_run_font(a, size=11, bold=True)
        b = p.add_run(text[len(bold_lead):])
        set_run_font(b, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.10
        r = p.add_run(item)
        set_run_font(r, size=10.8)


def new_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        p_style = abstract.find(".//" + qn("w:pStyle"))
        if p_style is not None and p_style.get(qn("w:val")) == "ListNumber":
            abstract_id = abstract.get(qn("w:abstractNumId"))
            break
    if abstract_id is None:
        abstract_id = "0"
    existing = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def add_numbers(doc, items):
    num_id = new_numbering_id(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
        num_pr.get_or_add_numId().set(qn("w:val"), str(num_id))
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.10
        r = p.add_run(item)
        set_run_font(r, size=10.8)


def add_callout(doc, label, text, fill=LIGHT_BLUE, color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    a = p.add_run(label + "  ")
    set_run_font(a, size=10.5, bold=True, color=color)
    b = p.add_run(text)
    set_run_font(b, size=10.5, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("WHY-REPAIR  |  项目进展与下一阶段规划")
        set_run_font(r, size=8.5, color=MUTED)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = fp.add_run("内部协作版本  ·  2026-08-18")
        set_run_font(r, size=8.5, color=MUTED)


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("项目进展与下一阶段规划书")
    set_run_font(r, size=24, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("依赖图驱动的双 Agent 自然语言数学证明审计与局部修复 Harness")
    set_run_font(r, size=13.5, color=MUTED)

    add_table(doc, ["字段", "内容"], [
        ("版本", "v1.0 · 基于 2026-08-18 仓库状态"),
        ("用途", "新协作者入组、任务分配、能力边界与下一阶段执行"),
        ("事实来源", "Git 历史、PROJECT_INDEX.md、ROADMAP.md、里程碑文档、机器清单与测试"),
        ("核心结论", "工程闭环已形成；正式独立实验、外部证据与论文级验收尚未闭合"),
    ], [1800, 7560])

    add_callout(doc, "当前判断", "项目已经从概念方案推进到 M0–M8 的完整工程骨架，并完成 50 题交互式诊断、反例、修复和九方法评测链路；但它仍是工程与研究候选，而不是已经完成独立盲测并可直接宣称性能优势的论文结果。")

    add_heading(doc, "1. 项目目标与边界", 1)
    add_body(doc, "项目目标是审计已有自然语言数学证明，而不是自主求解开放问题。系统将证明拆成节点和依赖图，由 Evaluator 定位第一处失败并生成结构化错误或反例证书，由 Repair Generator 提交最小局部补丁，再由独立 Evaluator 复核，Controller 负责版本、撤销、回滚、重试和可复现账本。")
    add_heading(doc, "1.1 可以合理追求的成果", 2)
    add_bullets(doc, [
        "可追踪的自然语言证明节点、依赖边、局部义务和第一处错误定位。",
        "结构化 ErrorCertificate、CounterexampleCertificate、PatchProposal 与 PatchReview。",
        "假命题识别、局部补丁、后代失效和多轮连续修复。",
        "统一的基线、消融、预算、运行账本、失败计分和复现清单。",
        "面向代数及后续受控扩展领域的专家可审 benchmark 和工程系统。",
    ])
    add_heading(doc, "1.2 当前不能声称的成果", 2)
    add_bullets(doc, [
        "不能声称自然语言 Agent 的接受等价于 Lean/Coq 等形式证明。",
        "不能把未找到反例当成命题正确的证明。",
        "不能把交互式投影或共享历史预测称为九组独立模型实验。",
        "不能声称 Full System 已显著优于基线，除非完成独立 Provider 运行和预注册统计。",
        "不能伪造独立审核身份、签名、模型快照、response ID、token、延迟、成本或账单证据。",
        "不能把用户允许继续工程执行解释为科学结论门已经通过。",
    ])

    add_heading(doc, "2. 当前系统已经具备的闭环", 1)
    add_numbers(doc, [
        "M2/M3：从 50 题 Gold 与模型输出形成证明级和节点级诊断。",
        "M4：对局部或全局反例执行 A→B→A 的上下文冻结、精确执行和复核闭环。",
        "M5：针对当前节点版本生成最小补丁，独立审核后事务式应用，并使全部受影响后代失效。",
        "M5 v0.2：修复首错后若后代暴露新错误，继续为新的首错生成证书，而不是过早判定不可修复。",
        "M6：将九种方法、50 个样本、预算、账本和指标统一到可复现评测框架。",
        "M7/M8：已有正式候选数据治理、运行门、错误分析、论文和发布审查骨架。",
    ])

    add_heading(doc, "3. M0–M8 真实进度总览", 1)
    rows = [
        ("M0", "范围、术语、研究问题、负向边界", "工程契约已冻结", "严格双盲独立性无法追溯；如论文要求需前瞻重做"),
        ("M1", "共享 Schema、状态机、版本与 Controller", "工程完成", "新一轮真实 A/B 签署视发布要求补充"),
        ("M2", "50 题 pilot、双人标注、裁决与 Gold", "已完成", "来源、盲态和外部可发布证据仍需按正式标准强化"),
        ("M3", "Evaluator v1、50 题运行、指标与集成", "工程完成", "历史运行非盲；分割 Gold、隔离模块运行和 Provider 证据不足"),
        ("M4", "反例证书、精确核验、A/B/Controller", "工程完成", "双外部复核、语义忠实性和前瞻盲测未完成"),
        ("M5", "Repair、版本、回滚、后代重验、多轮修复", "交互式工程完成", "正式 Provider pilot、独立补丁复核、成本与外部代码审查未完成"),
        ("M6", "九方法基线/消融、预算、账本、评分", "交互式工程完成", "九方法未独立运行；正式显著性和科学比较未完成"),
        ("M7", "50 题交互式扩展、OPC-250 候选治理", "部分完成", "独立 A/B Gold、裁决、入口门和 Provider 原始记录阻塞"),
        ("M8", "错误分析、论文与发布门骨架", "工程审查候选", "论文核心主张、外部复现、发布权利和正式 release 均未完成"),
    ]
    add_table(doc, ["阶段", "完成内容", "当前状态", "尚缺"], rows, [720, 2880, 1800, 3960], status_col=2)

    add_heading(doc, "4. 已经取得的可量化结果", 1)
    add_heading(doc, "4.1 数据与诊断", 2)
    add_bullets(doc, [
        "M2 已形成 50 题 pilot Gold；M3 已形成完整诊断与节点/依赖评估产物。",
        "M4 对冻结 benchmark 中 11 个有效反例完成工程批量验收。",
        "这些结果可作为工程回归证据，但历史 Gold 暴露与盲态限制必须保留。",
    ])
    add_heading(doc, "4.2 M5 修复处置", 2)
    add_table(doc, ["项目", "数量/结果", "解释"], [
        ("来源案例", "50", "完整 pilot 范围"),
        ("无需补丁的有效案例", "14", "保持原证明"),
        ("需要 M5 处置", "36", "均有人工决策与完成记录"),
        ("接受修复", "24", "交互式工程范围内完成审核和重验"),
        ("原命题假或无定义", "12", "正确终止为不可修复"),
        ("交互式 false repair / 新错误", "0 / 0", "仅适用于现有审核材料，不代表未知数据泛化"),
    ], [2160, 1800, 5400])
    add_heading(doc, "4.3 M6 交互式工程指标", 2)
    add_table(doc, ["指标", "结果", "证据边界"], [
        ("首错精确定位", "27/37 = 72.97%", "来自共享历史预测"),
        ("无首错位置假阳性", "0/12", "工程投影"),
        ("错误证明误接受", "0/26", "工程投影"),
        ("假命题检测", "11/11", "依赖已有诊断/反例处置"),
        ("证明弃权", "1/50 = 2%", "与基础设施失败分开"),
        ("已审核修复成功", "24/24", "消费 M5 人工审核结果"),
        ("false repair / 新错误", "0/24 / 0/24", "不是新数据上的独立模型实验"),
    ], [2760, 2160, 4440])
    add_callout(doc, "解释限制", "M6 的 50×9=450 个终态验证了配置、账本、评分和失败语义，但九种方法共享底层历史预测。因此不能用当前数字计算或宣称 Full System 对基线的真实提升。", fill="FFF4D6", color=GOLD)

    add_heading(doc, "5. 当前真正能做与暂时不能做", 1)
    add_table(doc, ["类别", "现在可以做", "现在不能可靠做"], [
        ("工程开发", "继续改进 Controller、Schema、测试、数据校验、运行器和报告", "原地修改冻结资产或为通过测试重写历史证据"),
        ("数学审计", "在现有协议下做节点、首错、反例和补丁复核", "把模型一致意见当作绝对数学证明"),
        ("M5", "扩展交互式案例、改进连续修复、准备真实 pilot", "声称已有生产模型成功率、真实成本或外部独立验收"),
        ("M6", "运行 fixture、检查预算/缓存/分母、准备 Provider runner", "基于共享预测报告显著性或消融因果效果"),
        ("M7", "完善 OPC-250 候选、许可、去重、Gold 流程", "在独立 Gold 与入口门未闭合前启动正式主实验"),
        ("M8", "写方法、系统设计、限制和工程结果草稿", "提交带有未验证性能主张的最终论文或正式 release"),
    ], [1320, 3960, 4080])

    add_heading(doc, "6. 下一阶段优先计划", 1)
    add_heading(doc, "优先级 P0：统一事实与修复状态漂移", 2)
    add_bullets(doc, [
        "执行全仓库只读状态审计，列出 README、ROADMAP、PROJECT_INDEX、机器 manifest 之间的冲突。",
        "冻结一个新的“当前状态快照”，不重写历史版本；对已知哈希漂移用新 manifest 或勘误处理。",
        "为 M5/M6/M7 分别建立工程完成、执行放行、科学验收三个独立状态字段。",
    ])
    add_heading(doc, "优先级 P1：补齐真实 M5/M6 运行链", 2)
    add_bullets(doc, [
        "确定真实 Provider、模型快照、采样参数、统一截断器和价格表。",
        "运行小规模 M5 Repair Generator pilot，保存每次 attempt、response ID、token、延迟、重试、失败和成本。",
        "由独立人员复核成功补丁、false repair、新错误和问题改变；身份豁免只影响签名要求，不替代数学复核。",
        "锁定 M6 九方法配置并先做 smoke；通过后再执行独立 9×N 运行，不复用跨方法缓存。",
    ])
    add_heading(doc, "优先级 P2：正式 Gold 与实验", 2)
    add_bullets(doc, [
        "完成 OPC-250 候选的独立 A/B 标注、分歧裁决、全量 Gold 和泄漏审计。",
        "在看正式结果前冻结 RQ、主指标、预算、比较族、功效输入和统计方案。",
        "运行同模型双角色与至少一组异模型组合；保留全部失败样本并按 intention-to-treat 计分。",
        "完成 bootstrap 区间、配对随机化检验与 Holm 校正，再决定哪些性能主张成立。",
    ])
    add_heading(doc, "优先级 P3：论文与发布", 2)
    add_bullets(doc, [
        "先写不会因实验结果变化而失效的方法、系统、数据流程和限制章节。",
        "仅将通过正式证据门的指标写入摘要、主表和结论。",
        "完成外部复现、数据许可、隐私/密钥扫描、系统卡、复现命令和发布 manifest。",
    ])

    add_heading(doc, "7. 新同学建议分工与 30/60/90 天计划", 1)
    add_heading(doc, "推荐角色：执行复现与证据完整性协作者", 2)
    add_body(doc, "新同学不宜一开始修改数学 Gold 或共享 Schema。最合适的切入点是独立复核工程证据、运行器、账本、缓存和状态一致性，同时逐步熟悉数学协议。")
    add_table(doc, ["时间", "任务", "交付与验收"], [
        ("第 1–7 天", "只读熟悉与状态审计", "M0–M8 状态表、权威文件图、冲突清单；不修改代码"),
        ("第 2–4 周", "复现现有测试和 fixture", "独立环境复现记录、失败归因、哈希/依赖问题清单"),
        ("第 2 个月", "负责 M5/M6 Provider runner 与运行证据", "smoke、attempt ledger、token/延迟/成本、失败保留和缓存隔离"),
        ("第 3 个月", "协助 OPC-250 与正式实验", "数据完整性检查、实验矩阵、聚合复算和复现包"),
    ], [1320, 3360, 4680])

    add_heading(doc, "8. 协作与变更控制", 1)
    add_bullets(doc, [
        "共享 Schema、Gold、指标定义和正式实验配置必须共同审查。",
        "冻结文件不可原地修改；新增版本必须有迁移说明、勘误或新 manifest。",
        "数学接受权、补丁生成权和 Controller 执行权必须分离。",
        "每个 PR 同时报告结果、文件、测试、契约影响、数学假设、证据强度和限制。",
        "失败样本、重试、超时和无输出必须保留；不得只汇报成功运行。",
        "任何无法从仓库或外部原始记录验证的事实，都必须标为 unknown、pending 或 unavailable。",
    ])

    add_heading(doc, "9. 新协作者 AI 使用方式", 1)
    add_body(doc, "完整可复制提示词已保存于 prompts/collaborator_onboarding_prompt_2026-08-18.md，包含首次接手、具体开发、只读进度审计、数学/Gold 审查、代码/PR 复核和日常续接六种模板。")
    add_callout(doc, "首次会话", "使用“提示词 A”，要求 AI 只读审计并等待任务。不要一开始让它修改代码、重构目录或更新状态。")
    add_callout(doc, "日常任务", "使用“提示词 B/F”，必须填写具体目标和验收标准。涉及证明或 Gold 时使用提示词 D；涉及合并审查时使用提示词 E。")

    add_heading(doc, "10. 完成标准与决策门", 1)
    add_body(doc, "以下四层状态必须分别记录，不能相互替代：")
    add_numbers(doc, [
        "实现完成：代码和 Schema 已存在。",
        "工程验证完成：正反测试、fixture、回放和哈希清单通过。",
        "执行证据完成：真实模型、原始调用、预算、失败和成本记录齐全。",
        "科学/发布验收完成：独立 Gold、盲测、预注册统计、外部复现和权利审查通过。",
    ])
    add_callout(doc, "当前项目位置", "M0–M6 多数模块已达到第 2 层；M5/M6 的交互式材料部分覆盖第 3 层的流程形状，但缺少正式 Provider 原始证据；M7/M8 尚未达到第 4 层。", fill="FFF4D6", color=GOLD)

    add_heading(doc, "附录 A：新同学入组检查清单", 1)
    add_bullets(doc, [
        "已克隆正确 GitHub 仓库并记录当前 commit。",
        "已阅读 AGENTS.md、README、PROJECT_INDEX、ROADMAP 和两人工作计划。",
        "能解释 Evaluator、Repair Generator、Controller 的权限边界。",
        "能区分 accepted、invalid、undetermined、stale、irreparable。",
        "能说明修改节点为何必须使后代失效。",
        "能说明 M6 的 450 条终态为何不能证明九方法性能差异。",
        "能指出 M5/M6/M7 当前缺少的真实外部证据。",
        "第一次提交不修改冻结 Gold、历史 manifest 或共享 Schema。",
        "提交包含正反测试、实际输出和限制说明。",
    ])

    add_heading(doc, "附录 B：权威入口", 1)
    add_table(doc, ["入口", "用途"], [
        ("PROJECT_INDEX.md", "项目产物导航和状态索引"),
        ("ROADMAP.md", "M0–M8 当前里程碑状态和退出条件"),
        ("docs/project_validation_and_acceptance_plan.md", "全部验证角色、证据门和验收要求"),
        ("docs/two_person_work_plan.md", "成员 A/B 分工与协作协议"),
        ("docs/milestones/", "每个阶段的权威设计、验收与限制"),
        ("data/benchmarks/ 与 data/governance/", "机器清单、运行产物、审计与授权边界"),
        ("harness/、schemas/、tests/", "实现、契约和可执行验证"),
        ("prompts/collaborator_onboarding_prompt_2026-08-18.md", "新协作者 AI 的六套提示词"),
    ], [3600, 5760])

    add_header_footer(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
