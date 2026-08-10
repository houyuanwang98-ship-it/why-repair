from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path("docs/dual_agent_natural_language_proof_harness_proposal.docx")
BLUE = "2E5D7B"
DARK = "183447"
LIGHT = "EEF4F7"
GRAY = "666666"
GOLD = "9A6A12"


def font(run, size=11, bold=False, italic=False, color="000000"):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def fix_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_p(doc, text="", bold_prefix=None, italic=False, align=None, after=7):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.3
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        font(r2, italic=italic)
    else:
        r = p.add_run(text)
        font(r, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text))


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    font(r, size={1: 16, 2: 13, 3: 11.5}[level], bold=True,
         color=BLUE if level < 3 else DARK)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    fix_table_geometry(table, [9360])
    c = table.cell(0, 0)
    shade(c, LIGHT)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(label + "  "), bold=True, color=DARK)
    font(p.add_run(text), color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    fix_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        shade(table.rows[0].cells[i], "E4ECF1")
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(h), size=9.5, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            font(p.add_run(str(value)), size=9.2)
    fix_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.85)
sec.bottom_margin = Inches(0.85)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.42)
sec.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.3
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, DARK, 8, 4),
]:
    st = styles[name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

# Running furniture
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(hp.add_run("自然语言数学证明审计 Harness｜研究方案草案"), size=8.5, color=GRAY)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("讨论总结 · 2026年8月"), size=8.5, color=GRAY)

# Cover
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("研究方案"), size=12, bold=True, color=GOLD)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
font(p.add_run("依赖图驱动的双 Agent\n自然语言数学证明审计 Harness"), size=27, bold=True, color=DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
font(p.add_run("面向节点验证、错误定位、反例发现与局部修复的无需训练系统"), size=14, color=BLUE)
add_callout(doc, "核心设想", "Evaluator 生成可检验的错误证书；Repair Generator 只针对失败节点提交最小补丁；Evaluator 随后撤销并重验所有受影响节点。")
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("项目讨论总结与实施蓝图"), size=11, bold=True, color=GRAY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("版本 0.1｜2026年8月10日"), size=10, color=GRAY)
doc.add_page_break()

add_heading(doc, "执行摘要", 1)
add_p(doc, "本项目总体可行，适合作为一个以系统设计、验证协议、数据结构和实验评估为核心的研究型工程。它不训练新的基础模型，也不依赖 Lean 等形式化语言，而是调用现有大语言模型，通过确定性的 harness 管理两个职责非对称的 Agent，对自然语言数学证明进行逐节点审计。")
add_p(doc, "系统的核心竞争力不是简单地使用两个 Agent，而是建立一种受约束的通信与验收协议：Evaluator 必须在证明依赖图上定位失败推理边，给出错误类型、缺失条件或经核验的反例；Repair Generator 只能针对这一局部义务提交最小补丁；Evaluator 独立复核补丁，并使所有依赖旧节点的后代结果失效后重新验证。")
add_callout(doc, "结论", "双 Agent 路线应当保留并作为论文主线；论文创新性需要落在“错误证书、依赖感知、反例核验、局部补丁、后代撤销与重验证”这一完整闭环，而不是 Agent 数量本身。")

add_heading(doc, "1. 项目定位与研究边界", 1)
add_heading(doc, "1.1 项目目标", 2)
add_p(doc, "构建一个可追踪、可重复评测的自然语言证明审计 harness：将证明分解为节点和依赖图，逐节点寻找支持、漏洞与反例，并通过 Evaluator—Repair Generator 闭环尝试受控的局部修复。")
add_heading(doc, "1.2 明确的非目标", 2)
for x in [
    "不从零训练一个新的数学大模型。",
    "第一阶段不追求自主解决开放数学问题。",
    "不把自然语言证明自动转换为 Lean、Coq 等形式证明作为必要条件。",
    "不声称自然语言模型的接受结果等价于机器可证明的绝对正确。",
    "不以整篇证明重写作为默认修复方式。",
]: add_bullet(doc, x)
add_heading(doc, "1.3 系统输出的证据强度", 2)
add_table(doc, ["等级", "含义", "应如何表述"], [
    ("强负面证据", "找到满足全部相关假设且否定目标结论的具体反例", "已发现错误"),
    ("结构化支持", "局部推导闭合，或必要定理的条件得到核验", "在给定上下文中得到支持"),
    ("不确定", "既未闭合推导，也未找到有效反例", "未确定；不得当作正确"),
], [1500, 4700, 3160])

add_heading(doc, "2. 核心架构：非对称双 Agent + 确定性控制器", 1)
add_p(doc, "系统在概念上只有两个数学 Agent，但工程上由一个非模型控制器编排。控制器不是第三个推理 Agent，而是负责状态、格式、版本、缓存、重试、回滚和运行记录的普通程序。")
add_callout(doc, "闭环", "Controller → Evaluator → Error Certificate → Repair Generator → Patch → Evaluator Recheck → Accept / Reject / Undetermined")
add_heading(doc, "2.1 Evaluator Agent", 2)
for x in [
    "切分并分类证明节点。",
    "构建每个节点的直接依赖图，并生成自包含的节点命题。",
    "将“全局假设 + 直接依赖 ⊢ 当前结论”构造成局部证明义务。",
    "同时寻找支持证据与反驳证据，主动尝试构造反例。",
    "定位具体失败推理边，而不是只给出模糊评分。",
    "决定接收、拒绝或保留不确定状态，并触发受影响后代的重验证。",
]: add_bullet(doc, x)
add_heading(doc, "2.2 Repair Generator Agent", 2)
for x in [
    "只接收被定位的错误节点、直接依赖、失败边、错误类型、已有反例和修改预算。",
    "输出结构化局部补丁，而不是重写整篇证明。",
    "明确补丁使用了哪些前提、规则和已有节点。",
    "不得自行宣称修复成功；是否接受只能由 Evaluator 决定。",
    "无法在原题条件下修复时，应返回不可修复，而不是偷偷增加假设。",
]: add_bullet(doc, x)
add_heading(doc, "2.3 Harness Controller", 2)
for x in [
    "校验 Agent 输出的 JSON Schema 与图结构约束。",
    "维护节点版本、依赖边、验证状态和调用轨迹。",
    "限制修复轮数，检测重复失败，并执行停止策略。",
    "节点改变后，撤销其所有直接或间接后代的旧验证结果。",
    "记录模型、提示词哈希、温度、token、成本和随机种子等实验元数据。",
]: add_bullet(doc, x)

add_heading(doc, "3. Evaluator 的分阶段工作流", 1)
add_p(doc, "虽然多个职责由同一个 Evaluator Agent 承担，但不应在一次自由文本调用中混合完成。每一阶段必须具有独立输入输出契约，并由控制器验证。")
steps = [
    "证明切分：把连续自然语言证明分解为具有明确源文本位置的节点。",
    "节点分类：区分 definition、assumption、claim、calculation、conclusion、citation 和 proof_strategy。",
    "依赖恢复：为每个节点找出直接的更早依赖，并解释依赖关系。",
    "图验证：程序检查节点覆盖、边方向、自环、重复边和 DAG 性质。",
    "局部义务检查：只使用原题假设和验证过的直接父节点。",
    "反例搜索：对不闭合节点主动寻找满足前提但否定结论的实例。",
    "错误诊断：输出第一处失败、失败边、错误类型和修复约束。",
]
for s in steps: add_number(doc, s)

add_heading(doc, "4. 错误证书：双 Agent 的正式通信接口", 1)
add_p(doc, "Evaluator 不应仅返回“第 N 步错误”。它必须生成一个可由 Repair Generator 消费、也可由程序和第三方复核的错误证书。建议字段如下：")
for x in [
    "failed_node：发生问题的节点及其当前版本。",
    "failed_edge：所用前提节点与目标结论之间的失败推理关系。",
    "local_context：当前节点合法可用的全部条件。",
    "error_type：缺失假设、定理误用、计算错误、局部假命题等。",
    "missing_condition：若适用，指出缺少的精确条件。",
    "counterexample：若存在，给出变量赋值和逐项核验。",
    "repair_constraints：允许的修改范围、禁止改变的内容和最大步骤预算。",
]: add_bullet(doc, x)
add_callout(doc, "核心原则", "Repair Generator 修复的是一个明确的局部证明义务，而不是根据模糊批评重新猜测整篇证明。")

add_heading(doc, "5. 反例协议", 1)
add_p(doc, "找反例是系统区别于普通 LLM critic 的关键能力。反例不能只是模型随口给出的例子，必须通过独立核验协议。")
for x in [
    "检查反例是否满足当前节点的全部局部前提。",
    "检查反例是否满足原题的全部全局假设。",
    "检查它是否确实否定目标结论。",
    "区分只否定当前局部节点，还是同时否定原定理。",
    "能计算的部分优先用 Python、SymPy、穷举或数值程序复核。",
    "没有找到反例绝不构成正确性证明。",
]: add_bullet(doc, x)
add_table(doc, ["反例类型", "判定范围", "系统动作"], [
    ("local counterexample", "否定当前节点，但不一定否定原定理", "标记 false_local_claim；允许尝试替换步骤"),
    ("global counterexample", "满足原题全部假设并否定最终结论", "标记 false_theorem；停止整题修复"),
    ("invalid candidate", "遗漏假设或未真正否定结论", "拒绝反例；保持待验证状态"),
], [1900, 3660, 3800])

add_heading(doc, "6. 修复协议与依赖失效", 1)
add_heading(doc, "6.1 允许的补丁操作", 2)
for x in [
    "insert_before：在错误节点前插入最小桥接步骤。",
    "replace：用一个或多个局部节点替换错误节点。",
    "delete：删除冗余、循环或错误节点。",
    "add_assumption：仅作为“改变原题”的显式建议，不计为原问题修复成功。",
    "mark_irreparable：确认在原题条件下无法修复。",
]: add_bullet(doc, x)
add_heading(doc, "6.2 修改后的重新验证", 2)
add_p(doc, "修复节点 N 后，不能只重新验证 N。控制器必须将所有依赖 N 的直接或间接后代标记为 stale，并按拓扑顺序重新验证。这相当于事实图中的 revocation 机制，可防止旧错误继续污染后续判断。")
add_heading(doc, "6.3 终止条件", 2)
for x in [
    "同一节点达到最大修复次数。",
    "连续两次产生等价错误或等价补丁。",
    "发现满足原题条件的全局反例。",
    "修复必须改变原定理或增加未授权假设。",
    "Evaluator 的独立复核长期冲突，系统应返回 undetermined。",
]: add_bullet(doc, x)

add_heading(doc, "7. 防止双 Agent 合谋与共同盲点", 1)
add_p(doc, "双 Agent 可能使用相同模型或相似训练数据，因此不能默认两者意见一致就代表正确。系统应通过信息隔离和对抗式复核降低共同盲点。")
for x in [
    "Evaluator 不读取 Repair Generator 的隐藏推理过程，只读取补丁和公开证据。",
    "Repair Generator 不接触 Evaluator 的完整裁决提示词。",
    "复核时重新构造局部义务，而不是沿用第一次判断结论。",
    "Evaluator 必须尝试反驳补丁，而不仅是寻找支持。",
    "高风险节点可以进行两次独立采样或异模型复核。",
    "结论冲突时应保留不确定，而不是用多数语言说服力代替数学证据。",
]: add_bullet(doc, x)

add_heading(doc, "8. 仓库结构与索引体系", 1)
add_p(doc, "仓库应采用“索引导航到真实成果”的组织方式。索引不复制所有内容，只记录模块状态、验收标准和实现位置。")
structure = """README.md
PROJECT_INDEX.md
ROADMAP.md
milestones/        # 每个小目标的目标、验收标准和状态
agents/            # evaluator 与 repair_generator 的提示词和契约
harness/           # 控制器、状态机、失效传播和运行清单
proof_graph/       # 图数据结构与验证逻辑
counterexamples/   # 反例生成、核验与模板
schemas/           # 统一 JSON Schema
benchmarks/        # 人工金标与压力测试数据
evaluations/       # 指标、基线与外部模型评估
experiments/       # 可复现实验配置和运行结果索引
docs/              # 设计说明、限制和论文材料
tests/             # 单元测试、契约测试和端到端测试"""
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.25)
p.paragraph_format.space_after = Pt(10)
for line in structure.splitlines():
    r = p.add_run(line + "\n")
    r.font.name = "Consolas"
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9.3)
    r.font.color.rgb = RGBColor.from_string(DARK)
add_p(doc, "PROJECT_INDEX.md 建议只保存模块名称、状态、规格链接、实现链接、测试链接和外部评估链接。每个 milestone 使用同一模板：目标、非目标、输入输出契约、验收标准、测试数据、实现位置、外部评估、已知限制和下一步。")

add_heading(doc, "9. 分阶段里程碑", 1)
milestones = [
    ("M0", "系统边界", "定义节点、依赖、正确、反例、修复成功和不确定", "术语与判定规范通过人工评审"),
    ("M1", "统一契约", "冻结 proof、node、edge、verdict、counterexample、patch schema", "全部示例通过 Schema 校验"),
    ("M2", "节点模块", "只做切分和分类", "在人工金标集上报告边界 F1 与分类 macro-F1"),
    ("M3", "依赖图", "只恢复直接依赖，不判断数学对错", "边 F1、DAG 合法率和关键依赖遗漏率"),
    ("M4", "局部验证", "使用人工正确节点和依赖隔离测试验证能力", "错误定位、假接受率和不确定校准"),
    ("M5", "反例模块", "构造并核验局部/全局反例", "发现率、有效率和假反例率"),
    ("M6", "Repair Generator", "在已知错误位置下生成最小补丁", "修复成功率和新错误引入率"),
    ("M7", "闭环 Harness", "加入版本、撤销、重验、重试和停止策略", "端到端状态机与回归测试通过"),
    ("M8", "外部评估", "多模型、人工专家和消融实验", "形成可复现论文实验包"),
]
add_table(doc, ["阶段", "主题", "主要产物", "验收重点"], milestones, [800, 1500, 3970, 3090])

add_heading(doc, "10. Benchmark 与评测设计", 1)
add_heading(doc, "10.1 数据构成", 2)
for x in [
    "正确证明与人工金标依赖图。",
    "从正确证明中注入单一、可控错误形成的配对样本。",
    "学生或模型自然产生的真实错误，避免只评估人工模板。",
    "跨错误类型数据：消去条件、量词、定理前提、计算、循环依赖、目标错配等。",
    "跨数学领域数据，但第一阶段可先聚焦代数以控制范围。",
]: add_bullet(doc, x)
add_heading(doc, "10.2 核心指标", 2)
add_table(doc, ["模块", "建议指标"], [
    ("节点", "Segmentation F1；Node classification macro-F1"),
    ("依赖", "Edge precision / recall / F1；关键依赖遗漏率"),
    ("验证", "First-error accuracy；verdict macro-F1；false acceptance rate"),
    ("反例", "Discovery rate；validity rate；assumption satisfaction rate"),
    ("修复", "Repair success；minimality；new-error introduction rate"),
    ("系统", "Descendant revalidation correctness；调用次数；token；成本；延迟"),
], [2200, 7160])
add_callout(doc, "优先指标", "数学验证场景中，错误证明被系统接受的比例通常比正确证明被暂时拒绝更危险，因此 false acceptance rate 应当是首要安全指标之一。")

add_heading(doc, "11. 双 Agent 核心价值的消融实验", 1)
add_p(doc, "论文必须证明性能来自双 Agent 的结构化闭环，而不仅仅是更多模型调用。建议比较：")
add_table(doc, ["方法", "图", "错误证书", "反例", "局部修复"], [
    ("单模型直接判断", "否", "否", "否", "否"),
    ("单模型自我反思", "否", "自由文本", "可选", "是"),
    ("普通 Generator–Critic", "否", "自由文本", "可选", "是"),
    ("双 Agent + 依赖图", "是", "是", "否", "是"),
    ("双 Agent + 反例", "否", "是", "是", "是"),
    ("完整系统", "是", "是", "是", "是"),
], [3100, 1100, 1800, 1400, 1960])
add_p(doc, "还应比较同模型双角色与异模型双角色、自由文本反馈与结构化错误证书、一次修复与多轮修复，以及是否执行后代撤销。若完整系统显著降低错误接受率并减少修复引入的新错误，双 Agent 才能成为有实验支撑的核心贡献。")

add_heading(doc, "12. 论文可行性与潜在贡献", 1)
add_p(doc, "项目具有论文潜力，但仅实现两个 Agent 的循环不足以构成强创新。更有说服力的论文定位是：一个无需训练、只使用自然语言、面向已有数学证明审计的依赖感知双 Agent harness。")
add_heading(doc, "12.1 推荐的研究问题", 2)
add_callout(doc, "Research Question", "错误证书驱动的双 Agent 闭环，是否比单模型判断、自我反思和普通 Generator–Critic 更可靠地定位、反驳并局部修复自然语言数学证明中的错误？")
add_heading(doc, "12.2 可主张的贡献", 2)
for x in [
    "定义 dependency-grounded natural-language proof auditing 新任务。",
    "提出错误证书驱动的非对称双 Agent 通信协议。",
    "提出反例核验、局部补丁及依赖后代撤销—重验证机制。",
    "发布带节点、依赖、错误位置、反例和修复金标的 benchmark。",
    "通过多模型和消融实验揭示自然语言证明验证的能力边界。",
]: add_bullet(doc, x)
add_heading(doc, "12.3 发表层级的现实判断", 2)
add_table(doc, ["完成度", "合理预期"], [
    ("仅原型和少量案例", "技术报告、课程项目或 demo；论文证据不足"),
    ("稳定系统 + 100–500 个高质量标注样本 + 消融", "适合 workshop 或 arXiv 技术报告"),
    ("多领域专家标注 benchmark + 强基线 + 系统性发现", "具备主会或专业会议投稿潜力"),
], [3300, 6060])

add_heading(doc, "13. 与相关系统的关系", 1)
add_p(doc, "Rethlas 已展示自然语言 generation agent 与 verification agent 的基本闭环；Danus 进一步以共享事实图、无状态验证器和撤销机制组织长期证明搜索。本项目保留这类 generate–verify–revise 思路和事实依赖管理，但主动缩小目标：不做开放问题的自主证明搜索，不做多 Worker 并行研究，不以 Lean 形式化为必要环节，而专注于已有自然语言证明的精细审计、反例发现和受控局部修复。")
add_p(doc, "参考链接：Rethlas，https://github.com/frenzymath/Rethlas；Danus: Orchestrating Mathematical Reasoning Agents with Fact-Graph Memory，https://arxiv.org/abs/2607.06447。")

add_heading(doc, "14. 算力与模型策略", 1)
add_p(doc, "第一阶段不需要训练，因此 A100 不是项目瓶颈。应优先投入高质量标注数据、稳定契约、状态管理和可复现实验。可以先用强 API 模型验证 harness；benchmark 稳定后，再使用 A100 部署开源模型、批量跑实验、比较同模型与异模型组合，并控制模型版本。")
for x in [
    "阶段一：强 API 模型，验证方法和数据契约。",
    "阶段二：A100 上部署开源模型，做规模化可复现实验。",
    "阶段三：只有在明确发现模型能力瓶颈且已有稳定数据后，才考虑微调。",
]: add_number(doc, x)

add_heading(doc, "15. 下一步建议", 1)
add_p(doc, "最稳妥的下一步不是立即扩充更多 Agent，而是先冻结第一个可评估版本的任务契约。建议按以下顺序执行：")
for x in [
    "撰写 M0 系统边界与术语规范，明确所有状态的判定含义。",
    "设计 proof、node、edge、error_certificate、counterexample 和 patch 的 JSON Schema。",
    "从当前仓库中抽取已有节点分类、依赖图、反例与失效逻辑，映射到新契约。",
    "建立首批 30–50 个专家可人工复核的代数样本，先验证评测流程。",
    "实现不调用模型也能测试的确定性 Controller 状态机。",
    "再接入 Evaluator 和 Repair Generator，逐模块形成基线与消融。",
]: add_number(doc, x)
add_callout(doc, "最终定位", "保留双 Agent，并将核心竞争力定义为：两个 Agent 之间传递可核验的错误证书；Evaluator 拥有裁决权；Repair Generator 只提交局部补丁；证明图负责管理真值状态、撤销与重新验证。")

# Keep headings with following content and avoid widows where possible.
for p in doc.paragraphs:
    p.paragraph_format.widow_control = True

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
